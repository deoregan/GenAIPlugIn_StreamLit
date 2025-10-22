# Import python libraries:
import os
import re
import pandas as pd
import jinja2
import pdfkit
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from bs4 import BeautifulSoup
import base64
import tempfile

# Import catalog and helpers:
from catalog.catalog import filenames
from catalog.llm_prompts import set_llm_prompts
from helpers.utils import image_to_base64, count_pass_fail
from helpers.formatting import format_numeric, conditional_formatting, highlight_rows, format_word_table
from helpers.generate_text import retrieve_quantity_from_table, wording_percentage_movement, wording_percentage_point_movement
from helpers.generate_text import wording_scr_movement, wording_target_solvency, wording_bscr_movements
from helpers.api_calls import llm_response

# Functions: 
#
#   create_html_report
#   html2pdf
#   crate_pdf_report
#   create_word_report
#   create_validation_report 

# Create html report:
def create_html_report(folders, filenames, current_year, previous_year, scr_table,   
                       llm_flag,  llm_provider, llm_nr_of_sentences,
                       target_solvency_ratio, conclusion_wording):

    """Function to create HTML report for SCR analysis.

    Args:
        folders (dict): Dictionary containing folder paths.
        filenames (dict): Dictionary containing file names.
        current_year (int): The current year.
        previous_year (int): The previous year.
        scr_table (pd.DataFrame): DataFrame containing SCR data.
        llm_flag (str): Flag indicating whether to use LLM or not.
        llm_provider (str): The LLM provider to use (e.g., 'OpenAI' or 'Gemini').
        llm_nr_of_sentences (int): Number of sentences to generate with LLM.
        target_solvency_ratio (float): Target solvency ratio.
        conclusion_wording (str): Conclusion wording for the report.

    Returns:
        str: The file path of the created HTML report.
    """

    # Set folders and filenames:
    root_folder                            = folders['root']
    output_images_folder                   = folders['output_images'] + str(current_year) + '/'
    bscr_current_chart_filename            = str(current_year) + '_' +  filenames['bscr_current_chart'] 
    bscr_previous_chart_filename           = str(previous_year) + '_' +  filenames['bscr_previous_chart']

    # Create HTML report:
    print(f"Creating HTML report... ")
    
    template_loader                        = jinja2.FileSystemLoader(searchpath="./")
    template_env                           = jinja2.Environment(loader=template_loader)
    layout_folder                          = folders['layout']
    report_layout_filename                 = filenames['scr_report_layout']
    template_file_reports                  = os.path.join(layout_folder, report_layout_filename)
    template                               = template_env.get_template(template_file_reports)

    # Generate the wording for the movements in the report:
    current_solvency_ratio                 = retrieve_quantity_from_table(scr_table, "Solvency Ratio", current_year)

    # Genarate wording for movements using Python code:
    scr_percentage_movement_wording_code   = wording_percentage_movement(scr_table, "Total SCR", previous_year, current_year, "m")
    scr_movement_wording_code              = wording_scr_movement(scr_table, previous_year, current_year, "m")
    own_funds_movement_wording_code        = wording_percentage_movement(scr_table, "Own Funds", previous_year, current_year, "m")
    solvency_ratio_movement_wording_code   = wording_percentage_point_movement(scr_table, "Solvency Ratio", previous_year, current_year)
    target_solvency_ratio_wording_code     = wording_target_solvency(target_solvency_ratio, current_solvency_ratio)
    bscr_movement_wording_code             = wording_bscr_movements(scr_table, 'amount', previous_year, current_year)


    # Set LLM prompts for AI wording:
    results_table                 = scr_table.to_string()
    prompts                       = set_llm_prompts(llm_nr_of_sentences, previous_year, results_table)  # XX used scr table not formatted one
    background_prompt             = prompts['background']
    results_analysis_prompt       = prompts['results_analysis']

    # Generate wording for the Background section (either with AI or llm_flag = False then with Python code only):
    background_wording            = llm_response(background_prompt, llm_flag, provider = llm_provider)

    # Generate results analysis (either with AI or llm_flag = False then with Python code only):   
    results_analysis_wording      = llm_response(results_analysis_prompt, llm_flag, provider = llm_provider)

    # Condition to use LLM response or code generated and reformat for inclusion in the table:
    if llm_flag == 'Yes':
        
        results_analysis_wording_reformat = results_analysis_wording

        #Remove redundant characters (: and -):
        results_analysis_wording_reformat = results_analysis_wording_reformat.replace(":", "").replace("-", "")           

        # If key items (SCR, Own Funds and Solvency Ratio) highlighed by double **:
        if '**' in results_analysis_wording_reformat:
            # Split the text into sentences
            sentences = re.split(r'(?<=[.!?])\s+', results_analysis_wording_reformat)
            results_analysis_wording = "<ul>"

            # Add bullets for each highlighted term (e.g., **SCR**, **Own Funds**, **Solvency Ratio**)
            for sentence in sentences:
                # Check for highlighted terms and add them as separate bullets
                if '**' in sentence:
                    # Extract term between '**' and add it as a bullet point
                    highlighted_term = re.findall(r"\*\*(.*?)\*\*", sentence)
                    if highlighted_term:
                        results_analysis_wording += f"<li>{highlighted_term[0]}</li>"  # Add term as a separate bullet

                    # Now add the rest of the sentence as another bullet point
                    sentence = re.sub(r"\*\*(.*?)\*\*", '', sentence)  # Remove the '**' part
                    if sentence.strip():
                        results_analysis_wording += f"<li>{sentence.strip()}</li>"
            
            results_analysis_wording += "</ul>"

        # Otherwise use simple bullet points: 
        else:

            print('No there are no **s')

            # Split sentences more reliably by ensuring there's a space after punctuation (or no space at all)
            sentences = re.split(r'(?<=[.!?])\s+', results_analysis_wording_reformat)  # Split by punctuation followed by space(s)

            results_analysis_wording = "<ul>"  # Start an unordered list
            for sentence in sentences:
                results_analysis_wording += f"<li>{sentence.strip()}</li>"  # Wrap each sentence in <li> tags
            results_analysis_wording += "</ul>"  # Close the unordered list

    else:
        results_analysis_wording = f"""
        <ul>
           <li>{scr_percentage_movement_wording_code}{scr_movement_wording_code}</li>
           <li>{own_funds_movement_wording_code}</li>
           <li>{solvency_ratio_movement_wording_code}</li>
           <li>{bscr_movement_wording_code}</li> 
        </ul>
          """

    # Generate wording for BSCR movements using Python code:
    bscr_percentage_movement_wording_code = wording_bscr_movements(scr_table, 'percentage', previous_year, current_year)

    # Create a copy of df for display purposes
    df_display = scr_table.copy()
    df_display.iloc[-1, 1:] = df_display.iloc[-1, 1:].apply(lambda x: f"{x:.1%}" if isinstance(x, (int, float)) and pd.notna(x) else x)

    # Convert images to base64 for display in the HTML report:
    bscr_current_chart_full_path = os.path.join(root_folder, output_images_folder, bscr_current_chart_filename)
    bscr_previous_chart_full_path = os.path.join(root_folder, output_images_folder, bscr_previous_chart_filename)
    
    bscr_current_chart_base64 = image_to_base64(bscr_current_chart_full_path)
    bscr_current_chart_html_tag = f'<img src="data:image/png;base64,{bscr_current_chart_base64}" alt="Image 1" style="width: 100%; height: auto;/">'

    bscr_previous_chart_base64 = image_to_base64(bscr_previous_chart_full_path)
    bscr_previous_chart_html_tag = f'<img src="data:image/png;base64,{bscr_previous_chart_base64}" alt="Image 1" style="width: 100%; height: auto;"/>'

    # Render html report:
    report_html     = template.render(df = df_display,
                                      current_year = current_year,
                                      previous_year = previous_year,
                                      background_wording = background_wording,  
                                      solvency_ratio_movement_wording_code = solvency_ratio_movement_wording_code,                           
                                      results_analysis_wording = results_analysis_wording,
                                      bscr_percentage_movement_wording_code = bscr_percentage_movement_wording_code,
                                      target_solvency_ratio_wording_code = target_solvency_ratio_wording_code,
                                      conclusion_wording = conclusion_wording,
                                      bscr_current_chart_html_tag = bscr_current_chart_html_tag,
                                      bscr_previous_chart_html_tag = bscr_previous_chart_html_tag

    )  

    # Set output reports folder:
    output_reports_folder = folders['output_reports'] + str(current_year) + '/'

    # Export HTML output file
    html_path = f'./{output_reports_folder}scr_report_{current_year}.html'

    with open(html_path, 'w', encoding="utf-8") as html_file:
        html_file.write(report_html)

        print(f"✅HTML SCR report saved at: {html_path}")
        
    # Save output path to dictionary:
    report_paths = {'html': html_path}

    return report_paths, report_html


def html2pdf(html_path, pdf_path):
    
    """
    Convert html to pdf using pdfkit.
    
    Args:
        html_path (str): Path to the input HTML file.
        pdf_path (str): Path to the output PDF file.
    Returns:
        None
    """
    
    options = {
        'page-size': 'A4',
        'margin-top': '0.35in',
        'margin-right': '0.75in',
        'margin-bottom': '0.75in',
        'margin-left': '0.75in',
        'encoding': "UTF-8",
        'no-outline': None,
        'enable-local-file-access': None
    }

    # Open the HTML file with UTF-8 encoding
    with open(html_path, 'r', encoding='utf-8') as f:
        pdfkit.from_file(f, pdf_path, options=options)



def create_pdf_report(folders, report_paths, current_year):
  
    """Function to create PDF report from HTML report.
    
    Args:
        folders (dict): Dictionary containing folder paths.
        report_paths (dict): Dictionary containing report paths.
        current_year (int): The current year for the report.

    Returns:
        dict: Updated report paths including the PDF path.
    """

    # Set folders:
    output_reports_folder = folders['output_reports'] + str(current_year) + '/'

    # Convert to pdf:
    print(f"Now converting to pdf... ")
        
    # Specify path:
    pdf_path_report = f'./{output_reports_folder}scr_report_{current_year}.pdf'  
        
    # Set HTML path:
    report_path_html = report_paths['html']

    # Create pdf file and confirm once ready:
    html2pdf(report_path_html, pdf_path_report)   
    print(f"✅PDF SCR report saved at: {pdf_path_report}.")

    # Add path to pdf file to outputs:
    report_paths['pdf'] = pdf_path_report

    return report_paths

def create_word_report(folders, report_paths, report_html, current_year):

    """Function to create Word report from HTML report.     

    Args:
        folders (dict): Dictionary containing folder paths.
        report_paths (dict): Dictionary containing report paths.
        report_html (str): HTML content of the report.
        current_year (int): The current year for the report.
        
    Returns:      
        dict: Updated report paths including the Word document path."""
    
    # Set folders:
    output_reports_folder = folders['output_reports'] + str(current_year) + '/'
    
    # Convert to Word:
    print(f"Now converting to Word (docx)... ")

    # Set path for saving the Word document
    docx_path = f'./{output_reports_folder}scr_report {current_year}.docx'

    # Parse the HTML content
    soup = BeautifulSoup(report_html, "html.parser")

    # Create a new Word document
    doc = Document()

    # Process all elements in order
    for tag in soup.find_all(["h1", "h2", "h3", "p", "table", "figure", "ul"]):  # Include <ul> here
    
        if tag.name == "h1" and "report-title" in tag.get("class", []):  # Handling h1 with class 'report-title'
            para = doc.add_paragraph()
            run = para.add_run(tag.get_text())
            run.bold = True
            run.font.size = Pt(16)  # Font size 16 for the report title
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-align the report title
            para.paragraph_format.space_before = Pt(6)  # Add 6-point space before

        elif tag.name == "h2":  # Specifically handling h2 headings
            para = doc.add_paragraph()
            run = para.add_run(tag.get_text())
            run.bold = True
            run.font.size = Pt(14)  # Font size 14 for h2 heading
            para.paragraph_format.space_before = Pt(6)  # Add 6-point space before

        elif tag.name.startswith("h"):  # Handle other headings (like h3, h4, etc.)
            para = doc.add_paragraph()
            run = para.add_run(tag.get_text())
            run.bold = True
            run.font.size = Pt(14) if tag.name == "h1" else Pt(12)  # Adjust font size based on the heading level
            para.paragraph_format.space_before = Pt(6) # Add 6-point space before

        elif tag.name == "p":
            para = doc.add_paragraph()  # Create the paragraph only once
            para.paragraph_format.space_after = Pt(0)  # Ensure no extra space after paragraphs

            for part in tag.contents:
                if isinstance(part, str):
                    # Replace <br><br> with ONE real line break
                    part = part.replace("\n", "")  # Removes excessive line breaks
                    run = para.add_run(part)
        
                elif part.name == "br":
                    para.add_run().add_break(WD_BREAK.LINE)  # Adds a true line break


        elif tag.name == "table":  # Handle tables
            caption_tag = tag.find("caption")
            if caption_tag:
                caption_para = doc.add_paragraph()
                caption_run = caption_para.add_run(caption_tag.get_text())
                caption_run.italic = True
                caption_run.font.size = Pt(10)
                caption_para.paragraph_format.space_before = Pt(6)

            rows = tag.find_all("tr")
            num_cols = len(rows[0].find_all(["th", "td"]))
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.style = 'Table Grid'

            # Fill in the table with data
            for row_idx, row in enumerate(rows):
                cells = row.find_all(["th", "td"])
                for col_idx, cell in enumerate(cells):
                    word_cell = table.cell(row_idx, col_idx)
                    word_cell.text = cell.get_text()

            format_word_table(table)   

        elif tag.name == "ul":  # Handling unordered list <ul>
            for li in tag.find_all("li"):  # Find all <li> items within the <ul>
                para = doc.add_paragraph(style='List Bullet')  # Applying 'List Bullet' style
                para.paragraph_format.space_before = Pt(0)  # Removes unwanted spacing above list
                para.paragraph_format.space_after = Pt(0)  # Removes unwanted spacing below list
                para.add_run(li.get_text())  # Add the text of each list item


        elif tag.name == "figure":  # Handle images and captions
            img_tag = tag.find("img")
            caption_tag = tag.find("figcaption")

            if img_tag and "src" in img_tag.attrs:
                img_src = img_tag["src"]

                if img_src.startswith("data:image"):  # Handle base64 images
                    base64_str = img_src.split(",")[1]  # Remove "data:image/png;base64,"
                    img_data = base64.b64decode(base64_str)

                    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_img:
                        temp_img.write(img_data)
                        temp_img_path = temp_img.name  # Get temporary file path

                    # Insert the image and resize it
                    doc.add_paragraph().add_run().add_picture(temp_img_path, width=Inches(3.5), height=Inches(3))  # Resize here

                    # Remove temp file after use
                    os.remove(temp_img_path)

            if caption_tag:
                caption_para = doc.add_paragraph()
                caption_run = caption_para.add_run(caption_tag.get_text())
                caption_run.italic = True
                caption_run.font.size = Pt(10)

        
    # Save as a Word document
    doc.save(docx_path)
    print(f"✅ Word SCR report saved at: {docx_path}.")    

    # Add path to pdf file to outputs:
    report_paths['docx'] = docx_path

    return report_paths

def create_validation_report(df_check, folders, current_year, previous_year, validation_threshold):

    """Function to create validation report for SCR analysis.

    Args:
        df_check (pd.DataFrame): DataFrame containing validation check results.
        folders (dict): Dictionary containing folder paths.
        current_year (int): The current year.
        previous_year (int): The previous year.
        validation_threshold (float): Threshold for validation checks.

    Returns:
        tuple: Paths to the HTML and PDF validation reports.
    """
        
    # Set folders:
    output_reports_folder = folders['output_reports'] + str(current_year) + '/'
    layout_folder = folders['layout']
    
    # Apply the counting function to each cell in the DataFrame
    results = df_check.applymap(lambda val: count_pass_fail(val, validation_threshold))

    # Count the number of passes and fails
    pass_count = (results == 'pass').sum().sum()
    fail_count = (results == 'fail').sum().sum()

    if fail_count == 0:
        validation_conclusion_wording = f'<b>All validation checks</b> passed which means that the Solvency Position of Smart Insurance Ltd \
                                         at year-end {current_year} is internal internally consistent. The validation threshold used is {validation_threshold}.'
    else:
        validation_conclusion_wording = f'The validation checks for the Solvency Position of Smart Insurance Ltd show <b>{fail_count} \
                                          failed tests</b>. Further investigation is needed. The validation threshold used is {validation_threshold}.'

    
    # Apply formatting and styling
    df_check_styled = (
        df_check
        .style
        .applymap(lambda val: conditional_formatting(val, validation_threshold))
        .format(format_numeric)
        .apply(highlight_rows, axis=1)  # Apply bold styling to specific rows
        .set_caption("Table 1 - Results of the validation checks for the Solvency Position of Smart Insurance Ltd")
        .hide(axis='index')
    )

    # Convert the styled DataFrame to an HTML string and export:
    html_path_validation       = f'./{output_reports_folder}validation_report_{current_year}.html'
    df_check_html = df_check_styled.to_html(escape=False)

    # Load layout template with Jinja2:
    template_loader = jinja2.FileSystemLoader(searchpath="./")
    template_env = jinja2.Environment(loader=template_loader)
    validation_layout_filename = filenames['validation_report_layout_filename']
    template_file_validation = os.path.join(layout_folder, validation_layout_filename)
    template = template_env.get_template(template_file_validation)

    # Render html validation report:
    html_report_validation = template.render(df = df_check_html,
                                             current_year = current_year,
                                             previous_year = previous_year,
                                             validation_conclusion_wording = validation_conclusion_wording
    )

    # Save html validation report:
    with open(html_path_validation, 'w', encoding="utf-8") as html_file:
        html_file.write(html_report_validation)

    # Specify path:
    pdf_path_validation = f'./{output_reports_folder}validation_report_{current_year}.pdf'  

    # Create pdf validation report and confirm once ready:
    html2pdf(html_path_validation, pdf_path_validation)   
    print(f"✅PDF Validation report saved at: {pdf_path_validation}.")

    return html_path_validation, pdf_path_validation