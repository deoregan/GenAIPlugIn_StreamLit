import pandas as pd
import re
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT 
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

# Functions:
#
#   In html (and same applied in pdf):
#
#   format_scr_table
#   replace_bold
#   create_bullet_points
#   conditinal_formatting
#   format_numeric
#   highlight_rows
# 
#   In Word (with docx)   


def format_scr_table(df, previous_year, current_year):

    """Function to format the SCR report table for better display.

    Args:
        df (pd.DataFrame): DataFrame containing the SCR report data.
        previous_year (int): Previous year for comparison.
        current_year (int): Current year for comparison.

    Returns:
        pd.DataFrame: Formatted DataFrame.          
    """

    # Set display options for better number formatting

    # Format numeric columns in pandas:
    pd.options.display.float_format = '{:,.0f}'.format  

    # Round all rows in the table (except for the Movement % column) to one decimal places 
    # except for the last row (solvency ratio) which is rounded to 3 decimal places:
    for col in ["Movement", current_year, previous_year]:
        df.loc[df.index[-1], col] = df[col].iloc[-1].round(3)
        df.loc[df.index[:-1], col] = df[col].iloc[:-1].round(1)

        # Format the 'Movement %' column as percentage with one decimal place:                  
    df["Movement %"]   = df["Movement %"].apply(lambda x: f"{100*x:.1f}%")
  
    return df

# Function to replace '**' with <strong> for bold in HTML
def replace_bold(text):
    """ Function to replace '**' with <strong> tags for bold formatting in HTML.        
    Args:
        text (str): Input text with '**' for bold.
    Returns:
        str: HTML formatted text with <strong> tags.
    """

    # Replace **term** with <strong>term</strong>
    text = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", text)
    return text

def create_bullet_points(text):

    """ Function to create HTML bullet points from text.

    Args:
        text (str): Input text to convert into bullet points.
    Returns:
        str: HTML formatted bullet points.
    """

    # Split sentences by punctuation with space after
    sentences = re.split(r'(?<=[.!?])\s+', text)
    html_output = "<ul>"

    # Loop through sentences and wrap them in <li> tags
    for sentence in sentences:
        sentence = sentence.strip()
        if sentence:  # Only add non-empty sentences
            html_output += f"<li>{sentence}</li>"
    
    html_output += "</ul>"
    return html_output 

def conditional_formatting(value, threshold):

    """Function to apply conditional formatting based on a threshold.   
    
    Args:
        value (float): Numeric value to evaluate.
        threshold (float): Threshold for formatting.        
    Returns:
        str: CSS style string for formatting.           
    """

    if isinstance(value, (int, float)):
        if value > threshold:
            return 'background-color: red; color: white'
        else:
            return 'background-color: green; color: white'
    return ''

# Define function to format numeric values to 3 decimals
def format_numeric(value, decimals=2):
    
    """Function to format numeric values to a specified number of decimal places.
    Args:
        value (float): Numeric value to format.
        decimals (int): Number of decimal places.   
    Returns:
        str: Formatted numeric value as string.
    """
    
    if isinstance(value, (int, float)):
        return f"{value:.{decimals}f}"
    return value

# Define function to apply bold styling to specific rows
def highlight_rows(row):
    
    """Function to apply bold styling to specific rows in a DataFrame.

    Args:
        row (pd.Series): The row to evaluate.
    Returns:
        list: A list of CSS styles for each cell in the row.
    """

    if row.name in [5, 8, 9, 10]:  # Row indices 6 and 9-11
        return ['font-weight: bold'] * len(row)
    return [''] * len(row)


def format_word_table(table):

    """Function to format a Word table using python-docx.   
    
    Args:
        table (docx.table.Table): The Word table to format.
    Returns:
        None
    """

    # Function to set background color of a cell
    def set_cell_background(cell, hex_color):
        
        """
        Sets the background color of a cell using Word's XML structure.

        Args:
            cell (docx.table.Cell): The cell to modify.
            hex_color (str): The background color in hex format (e.g., "FFFFFF" for white).

        Returns:
            None 
        """
        # Ensure the cell has a property container
        tc_pr = cell._tc.get_or_add_tcPr()
        
        # Create the shading element
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:val'), 'clear')  # Required attribute
        shading_elm.set(qn('w:color'), 'auto')  # Text color (can be 'auto')
        shading_elm.set(qn('w:fill'), hex_color)  # Background color in hex (without #)

        # Append the shading element to cell properties
        tc_pr.append(shading_elm)
    
    # Apply formatting to the header row:
    for cell in table.rows[0].cells:
        set_cell_background(cell, "000080")  # Blue background
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)  # White font color

    # Apply light fill color to specific rows (0-based index: row 7 and rows 10-12)
    target_rows = [6, 9, 10, 11]
    light_fill_color = "D9E1F2"  # Light blue color (you can change this)

    for row_idx in target_rows:
        for cell in table.rows[row_idx].cells:
            set_cell_background(cell, light_fill_color)

    # **Ensure tblPr (table properties) exists**
    tbl_pr = table._tbl.find(qn("w:tblPr"))  # Try to find existing tblPr
    if tbl_pr is None:  
        tbl_pr = OxmlElement("w:tblPr")  # Create new tblPr
        table._tbl.insert(0, tbl_pr)  # Insert at the beginning

    # **Expand table width**
    tbl_width = OxmlElement("w:tblW")
    tbl_width.set(qn("w:w"), "5000")  # Increase width (max is 16383)
    tbl_width.set(qn("w:type"), "pct")  # Set width as a percentage of the page
    tbl_pr.append(tbl_width)

    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            # Check if the cell is in the first column
            if col_idx == 0:
                # Left align for the first column
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                # right align other columns:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Center align the table
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    
    def set_cell_borders(cell, border_size="4"):  # "4" is a moderate border width
        
        """Function to set borders for a table cell.

        Args:
            cell (docx.table.Cell): The cell to modify.
            border_size (str): The size of the border (1-24).       
        Returns:
            None            
        """
    
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn("w:tcBorders"))
        
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)

        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")  # Single line border
            border.set(qn("w:sz"), border_size)  # Border thickness (1-24)
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")  # Black color
            tc_borders.append(border)

    # Apply borders to all cells
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell)
